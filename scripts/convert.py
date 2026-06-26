#!/usr/bin/env python3
"""
convert.py — Produce a Waabanong Weekly Update .docx from a markdown draft.

Uses python-docx to structurally modify a base document (May 25 accepted deliverable
or most recent local weekly-updates .docx), replacing text content at known positions
without breaking the grey table headers, Features boxes, logo, or footer.

Usage:
    python scripts/convert.py <week-monday> --image <path/to/roadmap.png>
    python scripts/convert.py 2026-06-22 --image Z:/path/to/gantt.png

    <week-monday>  Monday of the reported work week (YYYY-MM-DD)
    --image        Path to updated roadmap PNG (REQUIRED — no silent fallback)
    --base         Override base document path (default: auto-resolved)
    --dry-run      Parse draft and show what would be written, without saving

Outputs:
    outputs/weekly-updates/Waabanong Weekly Update - Week of <week-monday>.docx
"""

import sys
import re
import os
import argparse
import shutil
from pathlib import Path
from copy import deepcopy
from datetime import date, timedelta

try:
    import docx
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)

REPO_ROOT = Path(__file__).parent.parent
OUTPUTS_WEEKLY = REPO_ROOT / "outputs" / "weekly-updates"
OUTPUTS_PEGUIS = REPO_ROOT / "outputs" / "peguis-cfs"
Z_FALLBACK = Path(
    r"Z:\Shared\3_Client Projects\Peguis CFS\Projects"
    r"\PD25-1186-CO - Custom CMS\4. Controlling\02. Status Reports"
    r"\Waabanong Weekly Update 2026-05-25.docx"
)

# Known table identifiers (first cell of row 0)
TABLE_GENERAL = "Project Management & General"
TABLE_ITERATIONS = [
    "I1 - System Infrastructure",
    "I2 - Login & Authorization",
    "I2 -  Login & Authorization",  # alt spacing in some files
    "I3 - Client, Member & Family",
    "I4 - Intake Management",
    "I5 A - Case File Management",
    "I5 B",
    "I6 - Data Cleaning",
    "I7 - Financial Tracking",
    "I8 - My Dashboard",
]

# Iteration header text in the draft → table identifier prefix in docx
DRAFT_TO_TABLE_KEY = {
    "Project Management & General": "Project Management & General",
    "I1 - System Infrastructure & Environments": "I1 - System Infrastructure",
    "I2 - Login & Authorization": "I2 - Login & Authorization",
    "I3 - Client, Member & Family Data": "I3 - Client, Member & Family",
    "I4 - Intake Management & Automated Notifications": "I4 - Intake Management",
    "I5 A - Case File Management & Signs of Safety": "I5 A - Case File Management",
    "I5 B – Abuse Investigations": "I5 B",
    "I5 B - Abuse Investigations": "I5 B",
    "I6 - Data Cleaning and Migration": "I6 - Data Cleaning",
    "I7 - Financial Tracking & Exports for Sage": "I7 - Financial Tracking",
    "I8 - My Dashboard & Reporting": "I8 - My Dashboard",
}


# ---------------------------------------------------------------------------
# Markdown parsing
# ---------------------------------------------------------------------------

def parse_draft(draft_path: Path) -> dict:
    """Parse a markdown draft into structured content."""
    text = draft_path.read_text(encoding="utf-8")

    # Extract "Week of:" date line
    date_m = re.search(r'\*\*Week of:\*\*\s*(.+)', text)
    week_date_str = date_m.group(1).strip() if date_m else ""

    # Extract summary table rows
    summary = {}
    table_block = re.search(
        r'## Iteration Summary\s*\n(.*?)\n(?:---|\Z)', text, re.DOTALL
    )
    if table_block:
        for line in table_block.group(1).split('\n'):
            m = re.match(
                r'\|\s*(I[0-9].*?)\s*\|\s*(.*?)\s*\|\s*(.*?)\s*\|\s*(.*?)\s*\|', line
            )
            if m:
                name = m.group(1).strip()
                summary[name] = {
                    'status': m.group(2).strip(),
                    'phase': m.group(3).strip(),
                    'notes': m.group(4).strip(),
                }

    # Split into sections by "## " headers
    section_re = re.compile(r'^## (.+)$', re.MULTILINE)
    boundaries = [(m.start(), m.group(1)) for m in section_re.finditer(text)]
    sections = {}
    for i, (start, name) in enumerate(boundaries):
        end = boundaries[i + 1][0] if i + 1 < len(boundaries) else len(text)
        body = text[start:end]
        if name == "Iteration Summary":
            continue
        features_m = re.search(
            r'\*\*Features\*\*\s*\n(.*?)(?=\*\*Weekly Progress\*\*|\Z)',
            body, re.DOTALL
        )
        progress_m = re.search(
            r'\*\*Weekly Progress\*\*\s*\n(.*?)(?=\n---|\Z)',
            body, re.DOTALL
        )
        sections[name] = {
            'features': features_m.group(1).strip() if features_m else "",
            'progress': _parse_bullets(progress_m.group(1).strip() if progress_m else ""),
        }

    return {
        'week_date': week_date_str,
        'summary': summary,
        'sections': sections,
    }


_PR_REF_RE = re.compile(r'\s*\(PR #\d+\)', re.IGNORECASE)


def _strip_reference_artifacts(text: str) -> str:
    """Remove internal reference artifacts that must not appear in Word output.
    Currently strips PR references like '(PR #50)'. The markdown draft retains
    them for traceability; they are cleaned here before writing to .docx.
    """
    return _PR_REF_RE.sub('', text).rstrip()


def _parse_bullets(text: str) -> list:
    """Parse markdown bullets into list of (indent_level, text) tuples."""
    bullets = []
    for line in text.split('\n'):
        stripped_right = line.rstrip()
        if not stripped_right.strip():
            continue
        stripped = stripped_right.lstrip()
        if stripped.startswith('- '):
            indent = (len(stripped_right) - len(stripped)) // 2
            bullets.append((indent, _strip_reference_artifacts(stripped[2:])))
    return bullets


# ---------------------------------------------------------------------------
# Document structure navigation
# ---------------------------------------------------------------------------

def get_body_map(doc) -> list:
    """
    Return list of (tag, element) tuples for direct children of body.
    tag is one of: 'p', 'tbl', 'other'
    """
    result = []
    for el in doc.element.body:
        raw_tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        if raw_tag in ('p', 'tbl'):
            result.append((raw_tag, el))
        else:
            result.append(('other', el))
    return result


def get_element_text(el) -> str:
    """Extract all text from an XML element (paragraph or table)."""
    return ''.join(t.text or '' for t in el.iter(qn('w:t')))


def find_table_by_prefix(doc, prefix: str):
    """Find a table whose first row first cell starts with the given prefix."""
    for tbl in doc.tables:
        if not tbl.rows:
            continue
        first_cell_text = tbl.rows[0].cells[0].text.strip()
        if first_cell_text.startswith(prefix):
            return tbl
    return None


# ---------------------------------------------------------------------------
# Content updates
# ---------------------------------------------------------------------------

def update_date_lines(doc, week_monday_iso: str):
    """Update the date line paragraphs near the top of the document."""
    try:
        week_monday = date.fromisoformat(week_monday_iso)
    except ValueError:
        print(f"WARNING: Could not parse week_monday '{week_monday_iso}', skipping date update")
        return

    week_friday = week_monday + timedelta(days=4)
    delivery_date = week_monday + timedelta(days=7)  # following Monday

    # Format: "Week of May 18 to 22, 2026"
    # Use .day to avoid %-d which doesn't work on Windows
    if week_monday.month == week_friday.month:
        week_of_str = (
            f"Week of {week_monday.strftime('%B')} {week_monday.day} "
            f"to {week_friday.day}, {week_friday.year}"
        )
    else:
        week_of_str = (
            f"Week of {week_monday.strftime('%B')} {week_monday.day} "
            f"to {week_friday.strftime('%B')} {week_friday.day}, {week_friday.year}"
        )
    delivered_str = f"Delivered {delivery_date.strftime('%B')} {delivery_date.day}, {delivery_date.year}"

    # Update first few non-title paragraphs
    date_paras_found = 0
    for p in doc.paragraphs[:12]:
        text = p.text.strip()
        if not text or p.style.name == 'Heading No Numbers':
            continue
        if text.startswith('Delivered') or text.startswith('Week of') or date_paras_found > 0:
            if date_paras_found == 0:
                _set_para_text(p, delivered_str)
            elif date_paras_found == 1:
                _set_para_text(p, week_of_str)
            else:
                break
            date_paras_found += 1
            if date_paras_found >= 2:
                break

    if date_paras_found == 0:
        print("WARNING: Could not find date line paragraphs to update")


def update_summary_table(doc, summary: dict):
    """Update the Iteration Summary table with new Status/Phase/Notes values."""
    tbl = None
    for t in doc.tables:
        if t.rows and 'Iteration' in t.rows[0].cells[0].text:
            tbl = t
            break
    if tbl is None:
        print("WARNING: Summary table not found")
        return

    # Map summary keys to row indices by matching iteration name prefix
    for row in tbl.rows[1:]:  # skip header row
        iter_name_in_doc = row.cells[0].text.strip().split('\n')[0]
        matched_key = _match_summary_key(iter_name_in_doc, summary)
        if matched_key:
            data = summary[matched_key]
            _set_cell_text(row.cells[1], data['status'])
            _set_cell_text(row.cells[2], data['phase'])
            _set_cell_text(row.cells[3], data['notes'])


def _match_summary_key(doc_name: str, summary: dict) -> str:
    """Find the summary key that best matches a document iteration name."""
    doc_name_clean = doc_name.lower().replace('–', '-').replace('  ', ' ')
    for key in summary:
        key_clean = key.lower().replace('–', '-').replace('  ', ' ')
        # Try prefix match on the iteration number
        iter_prefix = re.match(r'(i[0-9]+ ?[ab]?)', key_clean)
        doc_prefix = re.match(r'(i[0-9]+ ?[ab]?)', doc_name_clean)
        if iter_prefix and doc_prefix and iter_prefix.group(1) == doc_prefix.group(1):
            return key
    return None


def update_iteration_sections(doc, sections: dict):
    """Update Features and Weekly Progress for each iteration section."""
    body_map = get_body_map(doc)

    for section_name, data in sections.items():
        table_key = DRAFT_TO_TABLE_KEY.get(section_name)
        if not table_key:
            continue

        # Find the table element
        tbl = find_table_by_prefix(doc, table_key)
        if tbl is None:
            print(f"WARNING: Table not found for section '{section_name}' (key='{table_key}')")
            continue

        # Update Features (row 1 of iteration table, skip General which has 1 row)
        if section_name != "Project Management & General" and len(tbl.rows) >= 2:
            _update_features_cell(tbl.rows[1].cells[0], data['features'])

        # Replace Weekly Progress bullets
        if data['progress']:
            _replace_weekly_progress(doc, tbl._element, data['progress'])


def _update_features_cell(cell, features_text: str):
    """Replace the Features content in an iteration table cell."""
    if not features_text:
        return

    # Parse features as lines, preserving structure
    lines = [l for l in features_text.split('\n') if l.strip()]

    # The cell contains paragraphs. First paragraph is typically "Features" label.
    # We want to keep it and replace the rest.
    paras = cell.paragraphs
    if not paras:
        return

    # Keep the "Features" heading paragraph, replace/add content paragraphs
    feature_para_idx = None
    for i, p in enumerate(paras):
        if 'Features' in p.text:
            feature_para_idx = i
            break

    if feature_para_idx is None:
        return

    # Remove all paragraphs after the Features heading
    for p in paras[feature_para_idx + 1:]:
        p._element.getparent().remove(p._element)

    # Re-read paragraphs after removal
    features_heading_el = cell.paragraphs[feature_para_idx]._element

    # Insert feature lines after the Features heading
    for line in reversed(lines):
        stripped = line.lstrip('- ').lstrip()
        indent = (len(line) - len(line.lstrip())) // 2
        new_p = _create_list_para(cell.paragraphs[0], stripped, indent)
        features_heading_el.addnext(new_p)


def _replace_weekly_progress(doc, table_el, bullets: list):
    """
    Replace Weekly Progress bullets after the given table element.
    Finds the 'Section' paragraph ('Weekly Progress') immediately following the table,
    then replaces all 'list1' paragraphs until the next table or significant gap.
    """
    body = doc.element.body
    body_children = list(body)

    # Find position of table_el
    try:
        tbl_idx = body_children.index(table_el)
    except ValueError:
        print(f"WARNING: Could not find table element in body")
        return

    # Find "Weekly Progress" paragraph after the table
    weekly_progress_el = None
    weekly_progress_idx = None
    for i in range(tbl_idx + 1, min(tbl_idx + 8, len(body_children))):
        el = body_children[i]
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        if tag == 'p':
            text = get_element_text(el).strip()
            if text == 'Weekly Progress':
                weekly_progress_el = el
                weekly_progress_idx = i
                break
        elif tag == 'tbl':
            break  # Hit next table before finding Weekly Progress

    if weekly_progress_el is None:
        print(f"WARNING: No 'Weekly Progress' paragraph found after table")
        return

    # Find the next table (or end of meaningful content) after Weekly Progress
    next_table_el = None
    for i in range(weekly_progress_idx + 1, len(body_children)):
        el = body_children[i]
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        if tag == 'tbl':
            next_table_el = el
            break

    # Remove existing bullet paragraphs between weekly_progress and next_table
    to_remove = []
    started = False
    for el in body_children[weekly_progress_idx + 1:]:
        if next_table_el is not None and el is next_table_el:
            break
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        if tag == 'p':
            to_remove.append(el)

    for el in to_remove:
        body.remove(el)

    # Find an example list1 paragraph to copy its formatting
    example_list1 = _find_example_para(doc, 'list1')
    example_section = _find_example_para(doc, 'Section')

    # Insert new bullets after weekly_progress_el
    anchor = weekly_progress_el
    for level, text in bullets:
        new_p = _create_para_copy(example_list1, text, level)
        anchor.addnext(new_p)
        anchor = new_p  # insert after previous bullet (maintains order)


def _find_example_para(doc, style_name: str):
    """Find first paragraph with given style as template."""
    for p in doc.paragraphs:
        if p.style.name == style_name:
            return p._element
    return None


def _create_para_copy(example_el, text: str, indent_level: int = 0):
    """Create a new paragraph by copying an example element and replacing text."""
    if example_el is not None:
        new_p = deepcopy(example_el)
        # Remove all runs
        for r in new_p.findall(qn('w:r')):
            new_p.remove(r)
        # Remove any hyperlink elements too
        for hl in new_p.findall(qn('w:hyperlink')):
            new_p.remove(hl)
    else:
        # Fallback: minimal paragraph
        new_p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), 'list1')
        pPr.append(pStyle)
        new_p.append(pPr)

    # Add indentation for sub-bullets
    if indent_level > 0:
        pPr = new_p.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            new_p.insert(0, pPr)
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        base = 360
        ind.set(qn('w:left'), str(base + indent_level * 360))

    # Add text run
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    if text != text.strip():
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    new_p.append(r)
    return new_p


def _create_list_para(example_para, text: str, indent_level: int = 0):
    """Create a new list paragraph from an example paragraph element."""
    return _create_para_copy(example_para._element if example_para else None, text, indent_level)


# ---------------------------------------------------------------------------
# Section spacing
# ---------------------------------------------------------------------------

def add_section_spacing(doc, space_pt: float = 18.0):
    """
    Add breathing room before each section table (General + I1–I8) by setting
    space_after on the paragraph immediately preceding the table. This is the
    standard Word approach — paragraph spacing controls the visual gap above tables.

    18pt (360 twips) is the default; the grey header boxes themselves are unchanged.
    """
    space_twips = str(int(space_pt * 20))
    body = list(doc.element.body)

    for i, el in enumerate(body):
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        if tag != 'tbl':
            continue

        # Skip the summary table (its first row contains "Iteration" and "Status")
        rows = el.findall('.//' + qn('w:tr'))
        if not rows:
            continue
        first_text = ''.join(t.text or '' for t in rows[0].iter(qn('w:t')))
        if 'Iteration' in first_text and 'Status' in first_text:
            continue

        # Find the paragraph immediately before this table and set space_after
        for j in range(i - 1, -1, -1):
            prev = body[j]
            prev_tag = prev.tag.split('}')[-1] if '}' in prev.tag else prev.tag
            if prev_tag == 'p':
                _set_space_after(prev, space_twips)
                break


def _set_space_after(para_el, space_twips: str):
    """Set w:spacing w:after on a paragraph XML element (in twips)."""
    pPr = para_el.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para_el.insert(0, pPr)
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:after'), space_twips)


# ---------------------------------------------------------------------------
# Image replacement
# ---------------------------------------------------------------------------

def replace_roadmap_image(doc, image_path: Path):
    """Replace the first embedded image in the document with a new PNG."""
    # Find image relationships
    image_rels = [
        r for r in doc.part.rels.values()
        if 'image' in r.reltype
    ]
    if not image_rels:
        print("WARNING: No images found in document — cannot replace roadmap")
        return

    # Replace the first image (index 0 is typically the roadmap)
    first_image_rel = image_rels[0]
    image_part = first_image_rel.target_part
    with open(image_path, 'rb') as f:
        image_part._blob = f.read()
    # content_type is read-only on ImagePart; PNG is already the correct type
    # for this workflow since base documents always use PNG roadmap images.
    print(f"  Roadmap image replaced from {image_path.name}")


# ---------------------------------------------------------------------------
# Cell text helpers
# ---------------------------------------------------------------------------

def _set_cell_text(cell, text: str):
    """Replace all text in a table cell with the given string, preserving style."""
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ''
    if cell.paragraphs:
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = text
        else:
            p.add_run(text)
    else:
        cell.add_paragraph(text)


def _set_para_text(para, text: str):
    """Replace paragraph text preserving runs (updates first run, clears rest)."""
    for r in para.runs:
        r.text = ''
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


# ---------------------------------------------------------------------------
# Base document resolution
# ---------------------------------------------------------------------------

def _date_from_filename(p: Path) -> str:
    """Extract YYYY-MM-DD from a 'Week of YYYY-MM-DD' filename, or '' if not found."""
    m = re.search(r'(\d{4}-\d{2}-\d{2})', p.stem)
    return m.group(1) if m else ''


def find_base_document(output_path: Path) -> Path:
    """
    Resolve the base document to use for conversion.
    Priority: most recently-dated .docx from outputs/peguis-cfs/ (sorted by filename date,
    not mtime, because Z: drive copies preserve source timestamps).
    Then outputs/weekly-updates/ (pandoc-generated — flat structure, used only as last resort).
    Then Z: drive fallback.

    NOTE: outputs/peguis-cfs/ takes priority because outputs/weekly-updates/
    files are pandoc-generated and have flat Heading2 structure, not the
    correct grey-table structure of accepted client deliverables.
    """
    # Prefer peguis-cfs files (correct grey table structure) — sort by filename date
    candidates = sorted(
        list(OUTPUTS_PEGUIS.glob("*.docx")),
        key=_date_from_filename,
        reverse=True
    )
    if candidates:
        return candidates[0]

    # Fall back to weekly-updates (pandoc outputs — wrong structure, but better than nothing)
    candidates = sorted(
        [p for p in OUTPUTS_WEEKLY.glob("*.docx") if p.resolve() != output_path.resolve()],
        key=_date_from_filename,
        reverse=True
    )
    if candidates:
        print("WARNING: Using pandoc-generated .docx as base -- grey table structure may not be preserved")
        return candidates[0]

    # Z: drive
    if Z_FALLBACK.exists():
        return Z_FALLBACK

    raise FileNotFoundError(
        "No base document found. Ensure outputs/peguis-cfs/ has a .docx or Z: drive is accessible."
    )


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument('week_monday', help='Monday of the reported work week (YYYY-MM-DD)')
    parser.add_argument('--image', help='Path to roadmap PNG (REQUIRED)', metavar='PATH')
    parser.add_argument('--base', help='Override base document path', metavar='PATH')
    parser.add_argument('--dry-run', action='store_true', help='Parse and report without writing')
    args = parser.parse_args()

    # Validate week-monday
    try:
        date.fromisoformat(args.week_monday)
    except ValueError:
        print(f"ERROR: week_monday must be YYYY-MM-DD, got '{args.week_monday}'", file=sys.stderr)
        sys.exit(1)

    # Hard fail if no image supplied
    if not args.image:
        print(
            "ERROR: --image is required. Supply a path to the updated roadmap PNG.\n"
            "  The roadmap image is NEVER carried forward silently — provide it or\n"
            "  open the output .docx and replace the image manually after conversion.",
            file=sys.stderr
        )
        sys.exit(2)

    image_path = Path(args.image)
    if not image_path.exists():
        print(f"ERROR: Image file not found: {image_path}", file=sys.stderr)
        sys.exit(1)

    draft_path = OUTPUTS_WEEKLY / f"{args.week_monday}-draft.md"
    output_path = OUTPUTS_WEEKLY / f"Waabanong Weekly Update - Week of {args.week_monday}.docx"

    if not draft_path.exists():
        print(f"ERROR: Draft not found: {draft_path}", file=sys.stderr)
        print(f"  Run /weekly-update {args.week_monday} to generate the draft first.")
        sys.exit(1)

    print(f"Parsing draft: {draft_path.name}")
    content = parse_draft(draft_path)

    if args.dry_run:
        print("\n=== DRY RUN — parsed content ===")
        print(f"Week date: {content['week_date']}")
        print(f"Summary rows: {len(content['summary'])}")
        print(f"Sections: {list(content['sections'].keys())}")
        for name, data in content['sections'].items():
            bullets = data['progress']
            print(f"  {name}: {len(bullets)} bullet(s)")
        return

    # Resolve base document
    if args.base:
        base_path = Path(args.base)
        if not base_path.exists():
            print(f"ERROR: Base document not found: {base_path}", file=sys.stderr)
            sys.exit(1)
    else:
        base_path = find_base_document(output_path)

    print(f"Base document: {base_path.name}")

    # Copy base to output path (we edit the copy)
    OUTPUTS_WEEKLY.mkdir(parents=True, exist_ok=True)
    shutil.copy2(base_path, output_path)
    print(f"Output: {output_path.name}")

    # Open and update
    doc = docx.Document(output_path)

    print("  Updating date lines...")
    update_date_lines(doc, args.week_monday)

    print("  Updating summary table...")
    update_summary_table(doc, content['summary'])

    print("  Updating iteration sections...")
    update_iteration_sections(doc, content['sections'])

    print("  Adding section spacing (18pt before each grey header)...")
    add_section_spacing(doc, space_pt=18.0)

    print("  Replacing roadmap image...")
    replace_roadmap_image(doc, image_path)

    doc.save(output_path)
    print(f"\n[OK] Written: {output_path}")
    print(f"\nNext step: run scripts/qa_check.py against the output, then review in Word.")


if __name__ == '__main__':
    main()
