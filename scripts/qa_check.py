#!/usr/bin/env python3
"""
qa_check.py — Structural QA for a Waabanong Weekly Update .docx.

Checks that the document has the correct table-based structure matching the
accepted client deliverable format, not just text presence.

Usage:
    python scripts/qa_check.py <output.docx> [--fixture <approved.docx>]

Exit code: 0 if all checks pass, 1 if any fail.
"""

import sys
import argparse
from pathlib import Path

try:
    import docx
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)

EXPECTED_GREY_TABLES = [
    "Project Management & General",
    "I1 - System Infrastructure",
    "I2 - Login",
    "I3 - Client",
    "I4 - Intake",
    "I5 A",
    "I5 B",
    "I6 - Data",
    "I7 - Financial",
    "I8 - My Dashboard",
]

EXPECTED_SUMMARY_ROWS = 10  # header + 9 iterations (I1 through I8 including I5 A and I5 B)
EXPECTED_SECTIONS = 9  # I1 through I8 + General


def check(label: str, passed: bool, detail: str = "") -> bool:
    status = "[PASS]" if passed else "[FAIL]"
    line = f"  {status}  {label}"
    if detail:
        line += f" — {detail}"
    print(line)
    return passed


def run_checks(doc_path: Path, fixture_path: Path = None) -> bool:
    print(f"\n=== QA: {doc_path.name} ===\n")

    if not doc_path.exists():
        print(f"ERROR: File not found: {doc_path}", file=sys.stderr)
        return False

    doc = docx.Document(str(doc_path))
    passed = []

    # --- Check 1: Title paragraph ---
    title_found = any(
        p.text.strip() == "Waabanong Weekly Update"
        for p in doc.paragraphs[:5]
    )
    passed.append(check("Title paragraph 'Waabanong Weekly Update'", title_found))

    # --- Check 2: Date line present ---
    date_found = any(
        'Week of' in p.text
        for p in doc.paragraphs[:10]
    )
    passed.append(check("'Week of' date line present", date_found))

    # --- Check 3: Summary table has correct row count ---
    summary_tbl = None
    for t in doc.tables:
        if t.rows and 'Iteration' in t.rows[0].cells[0].text:
            summary_tbl = t
            break
    if summary_tbl:
        row_count = len(summary_tbl.rows)
        passed.append(check(
            f"Summary table has {EXPECTED_SUMMARY_ROWS} rows (header + 9 iterations)",
            row_count == EXPECTED_SUMMARY_ROWS,
            f"found {row_count}"
        ))
    else:
        passed.append(check("Summary table found", False, "not found"))

    # --- Check 4: Grey-table structure for all iteration sections ---
    table_prefixes_found = []
    for t in doc.tables:
        if not t.rows:
            continue
        first_cell = t.rows[0].cells[0].text.strip()
        for prefix in EXPECTED_GREY_TABLES:
            if first_cell.startswith(prefix):
                table_prefixes_found.append(prefix)

    all_sections_present = all(p in table_prefixes_found for p in EXPECTED_GREY_TABLES)
    passed.append(check(
        "All 10 section tables present (grey table structure)",
        all_sections_present,
        f"found {len(table_prefixes_found)}/{len(EXPECTED_GREY_TABLES)}: " +
        (", ".join(p for p in EXPECTED_GREY_TABLES if p not in table_prefixes_found) or "all present")
    ))

    # --- Check 5: Each iteration table has 2 rows (header + features) ---
    feature_tables_ok = True
    missing = []
    for t in doc.tables:
        if not t.rows:
            continue
        first_cell = t.rows[0].cells[0].text.strip()
        for prefix in EXPECTED_GREY_TABLES:
            if first_cell.startswith(prefix) and prefix != "Project Management & General":
                if len(t.rows) < 2:
                    feature_tables_ok = False
                    missing.append(prefix)
    passed.append(check(
        "Iteration tables have Features row (2 rows: header + features)",
        feature_tables_ok,
        "missing: " + ", ".join(missing) if missing else ""
    ))

    # --- Check 6: 'Section' style 'Weekly Progress' paragraphs ---
    section_paras = [p for p in doc.paragraphs if p.style.name == 'Section']
    wp_count = sum(1 for p in section_paras if 'Weekly Progress' in p.text)
    passed.append(check(
        f"'Weekly Progress' Section-style paragraphs: expected {EXPECTED_SECTIONS}",
        wp_count >= EXPECTED_SECTIONS,
        f"found {wp_count}"
    ))

    # --- Check 7: 'list1' style bullets present ---
    list1_count = sum(1 for p in doc.paragraphs if p.style.name == 'list1' and p.text.strip())
    passed.append(check(
        "list1-style bullets present (>= 5)",
        list1_count >= 5,
        f"found {list1_count}"
    ))

    # --- Check 8: No warning placeholder in final output ---
    warning_found = any(
        'Roadmap image' in p.text or 'provide updated PNG' in p.text
        for p in doc.paragraphs[:5]
    )
    passed.append(check(
        "No roadmap warning placeholder in output",
        not warning_found,
        "Warning placeholder still present -- was the draft converted or just copied?" if warning_found else ""
    ))

    # --- Check 9: I5 A and I5 B are separate sections ---
    i5a_found = any(
        t.rows and t.rows[0].cells[0].text.strip().startswith("I5 A")
        for t in doc.tables
    )
    i5b_found = any(
        t.rows and t.rows[0].cells[0].text.strip().startswith("I5 B")
        for t in doc.tables
    )
    passed.append(check(
        "I5 A and I5 B are separate section tables",
        i5a_found and i5b_found,
        f"I5A={i5a_found}, I5B={i5b_found}"
    ))

    # --- Check 10: Image relationship present ---
    image_rels = [r for r in doc.part.rels.values() if 'image' in r.reltype]
    passed.append(check(
        "At least one embedded image (roadmap/Gantt)",
        len(image_rels) >= 1,
        f"found {len(image_rels)} image rel(s)"
    ))

    # --- Optional: structural comparison against fixture ---
    if fixture_path and fixture_path.exists():
        print(f"\n--- Structural comparison vs fixture: {fixture_path.name} ---")
        fixture = docx.Document(str(fixture_path))
        fixture_tables = len(fixture.tables)
        output_tables = len(doc.tables)
        passed.append(check(
            f"Table count matches fixture ({fixture_tables})",
            output_tables == fixture_tables,
            f"output has {output_tables}"
        ))
        fixture_bullets = sum(1 for p in fixture.paragraphs if p.style.name == 'list1' and p.text.strip())
        passed.append(check(
            "Bullet count comparable to fixture (within 50%)",
            list1_count >= fixture_bullets * 0.5,
            f"output={list1_count}, fixture={fixture_bullets}"
        ))

    # --- Summary ---
    total = len(passed)
    n_pass = sum(passed)
    print(f"\n{'='*50}")
    print(f"Result: {n_pass}/{total} checks passed")
    if n_pass == total:
        print("[PASS] All checks passed.\n")
    else:
        print(f"[FAIL] {total - n_pass} check(s) failed -- review before delivering.\n")

    return n_pass == total


def main():
    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument('docx_path', help='Path to the .docx to check')
    parser.add_argument('--fixture', help='Path to approved .docx for structural comparison', metavar='PATH')
    args = parser.parse_args()

    ok = run_checks(Path(args.docx_path), Path(args.fixture) if args.fixture else None)
    sys.exit(0 if ok else 1)


if __name__ == '__main__':
    main()
